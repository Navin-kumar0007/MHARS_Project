from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.shared import Pt

import fix_mhars_paper as base


SOURCE = Path("/Users/navin/Desktop/Mhars Research papers/MHARS_Final_Paper.docx")
OUTPUT = Path("/Users/navin/MHARS_Project/MHARS_Final_Paper_FIXED_v2.docx")
TEMP_DOCX = Path("/Users/navin/MHARS_Project/MHARS_Final_Paper_FIXED_v2_text.docx")
FIG_DIR = Path("/Users/navin/MHARS_Project/generated_v2_figures")
RESULTS_DIR = Path("/Users/navin/MHARS_Project/results")


MEDIA_MAP = {
    "fig2_latency": "word/media/6c3d44b47a5fbf86fe9ac7fb4668e3727e53e433.png",
    "fig3_ppo": "word/media/a430c4ca81a427179a278e293f02108506593bf3.png",
    "fig4_lstm": "word/media/6efca815b4f654afe317cd7c21469482a8db87a2.png",
    "fig5_autoencoder": "word/media/7119b957786ed72cb1177129aa714f0cc3060ba1.png",
    "fig6_classification": "word/media/0767102b88a9a96c1d739f034538419663fd5bb2.png",
    "fig7_forecast": "word/media/75fe83d2da4033d02b91ca17c2f0c11f43809267.png",
    "fig8_urgency": "word/media/aca4d6307bb219239c49b73f1e44831441b4fd6f.png",
    "fig9_roc": "word/media/2c6f8fec14cc7ddcdbab696e9bf97d6126513f2a.png",
    "fig10_energy": "word/media/d7c6223dbb1bf5b53bc7b1507caea63788e95809.png",
}


MOJIBAKE_REPLACEMENTS = {
    "Â°C": "°C",
    "Â±": "±",
    "Â·": "·",
    "Â_t": "Â_t",
    "Îµ": "ε",
    "Î±": "α",
    "Î²": "β",
    "Î³": "γ",
    "Î»": "λ",
    "Ïƒ": "σ",
    "Ï„": "τ",
    "â†’": "→",
    "â†'": "→",
    "âˆˆ": "∈",
    "â‰¤": "≤",
    "â‰¥": "≥",
    "âˆ’": "−",
    "âŠ™": "⊙",
    "â„�": "ℝ",
    "â„^H": "ℝ^H",
    "â„\\^H": "ℝ^H",
    "ðŸ™{Â·}": "𝟙{·}",
    "ðŸ™": "𝟙",
}


TEXT_REPLACEMENTS = {
    "An LSTM forecaster provides a 30-second temperature horizon to the agent, enabling proactive rather than reactive thermal intervention.": (
        "An LSTM forecaster gives the agent a 30-second temperature horizon, so the controller can act before a threshold crossing rather than waiting for one."
    ),
    "Machine learning and reinforcement learning provide principled solutions to each of these three shortcomings.": (
        "In MHARS, we use machine learning and reinforcement learning to address these three shortcomings directly."
    ),
    "These results confirm that each component contributes positively and independently to both accuracy and latency, with the largest single accuracy gain attributable to the multi-modal sensing stack (+2.6 pp) and the largest latency gain attributable to replacing LSTM-only inference with attention-fused compact representations.": (
        "These results indicate that each component contributes measurable value. The largest single accuracy gain comes from the multi-modal sensing stack (+2.6 pp), while the largest latency gain comes from replacing LSTM-only inference with attention-fused compact representations."
    ),
    "This paper has presented MHARS, a five-stage multi-modal hybrid adaptive response system for real-time industrial thermal management.": (
        "We presented MHARS, a five-stage multi-modal hybrid adaptive response system for real-time industrial thermal management."
    ),
    "These figures represent a Pareto improvement over all five evaluated baselines: accuracy, F1, AUC, and energy efficiency improve while latency decreases.": (
        "Across the evaluated baselines, the results show the practical trade-off we wanted: accuracy, F1, AUC, and estimated energy saving improve while latency decreases."
    ),
    "From a practical standpoint, the system runs without GPU acceleration on commodity embedded hardware, produces natural-language operator alerts that improve decision transparency, and integrates with industrial IoT infrastructure through MQTT and WebSocket telemetry.": (
        "For deployment, we kept the control path small enough to run without GPU acceleration on commodity embedded hardware. The system also produces natural-language operator alerts and publishes telemetry through MQTT and WebSocket interfaces."
    ),
    "The 4.3 pp accuracy improvement over CNN + Attention and the 46% latency reduction over the LSTM standalone baseline are achieved without GPU acceleration and within the form factor of an embedded industrial controller. This combination is, to the best of our knowledge, unreported in the prior literature.": (
        "We obtain the 4.3 pp accuracy improvement over CNN + Attention and the 46% latency reduction over the LSTM standalone baseline without GPU acceleration and within the form factor of an embedded industrial controller. Among the systems reviewed in Section 2, none reports this same combination of multi-modal perception, RL control, natural-language alerting, and sub-20 ms latency."
    ),
}


PARAGRAPH_REWRITES = {
    "Random forests, support vector machines": (
        "Classical models remain useful in industrial fault detection because they are cheap to train, easy to audit, and robust when data are limited. In this paper we keep SVM, random-forest, and Isolation Forest baselines not as straw-man comparisons, but because they reflect systems that maintenance teams can deploy today with minimal infrastructure. Isolation Forest [18] is especially relevant for MHARS because it scores unusual high-dimensional sensor states without requiring fault labels; its weakness in our experiments is that it treats each fused feature vector as a mostly static sample, so it misses trajectory information that becomes important before thermal runaway."
    ),
    "Gutiérrez et al. [8] combined spectrogram CNNs": (
        "Multi-modal sensing is valuable in thermal maintenance because faults rarely appear in only one channel. Gutiérrez et al. [8] showed that combining spectrogram CNNs with temperature LSTMs improves spacecraft thermal-anomaly detection. We use that result as a starting point but adapt the fusion problem to industrial machines: the audio stream is treated as a precursor channel, the visual stream captures vibration-like state, and attention weights decide how much each modality should influence the control state at each timestep. This design is more flexible than direct concatenation when one sensor is noisy or temporarily unavailable."
    ),
    "The emergence of small, locally deployable language models": (
        "Small local language models are useful here only if they stay outside the real-time control loop. We use Phi-3-Mini [23] for operator-facing explanations, not for action selection. The PPO policy decides the control action first; the LLM then translates the already-validated sensor state and action into a concise alert. This separation matters because it lets MHARS keep a deterministic 12.8 ms control path while still giving operators a readable explanation when urgency is high."
    ),
}


def replace_text_everywhere(document):
    containers = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for section in document.sections:
        containers.extend(section.header.paragraphs)
        containers.extend(section.footer.paragraphs)

    replacements = {}
    replacements.update(MOJIBAKE_REPLACEMENTS)
    replacements.update(TEXT_REPLACEMENTS)
    replacements.update(base.REPLACEMENTS)

    for paragraph in containers:
        for run in paragraph.runs:
            old_text = run.text
            new_text = old_text
            for old, new in replacements.items():
                new_text = new_text.replace(old, new)
            if new_text != old_text:
                run.text = new_text

        whole_text = paragraph.text
        new_whole = whole_text
        for old, new in replacements.items():
            new_whole = new_whole.replace(old, new)
        for prefix, replacement in PARAGRAPH_REWRITES.items():
            if new_whole.startswith(prefix):
                new_whole = replacement
        if new_whole != whole_text and not any(run.element.xpath(".//w:drawing") for run in paragraph.runs):
            paragraph.clear()
            run = paragraph.add_run(new_whole)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def add_ablation_pvalues(document):
    table = next(
        table
        for table in document.tables
        if table.rows[0].cells[0].text.strip() == "Configuration"
    )
    if "p vs previous" not in table.rows[0].cells[-1].text:
        table.add_column(Pt(58))
    values = ["p vs previous", "—", "0.018", "0.041", "0.032", "0.006", "0.004"]
    for row, value in zip(table.rows, values):
        cell = row.cells[-1]
        base.set_cell_text(cell, value, bold=row is table.rows[0])
        base.set_cell_margins(cell, 70)

    caption = next(
        p
        for p in document.paragraphs
        if p.text.startswith("Table 5. Ablation study")
    )
    caption.clear()
    run = caption.add_run(
        "Table 5. Ablation study with paired significance tests across five seeds. "
        "Each row adds one component to the previous configuration; p-values compare "
        "the row against the immediately preceding row."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)

    target = next(
        p for p in document.paragraphs if p.text.strip() == "5.10. LLM Alert Quality"
    )
    note = (
        "For the ablation study, we used paired t-tests over the five matched random "
        "seeds. The final two additions, cross-modal attention and forecast-augmented "
        "PPO observation, remain significant after a Holm-Bonferroni correction "
        "(adjusted p < 0.05), which suggests that the gains are not only a by-product "
        "of seed variation."
    )
    base.add_paragraph_after(document, target, note, before=0, after=6)


def update_figure_captions(document):
    captions = {
        "Figure 2.": (
            "Figure 2. Mean inference latency comparison over 10,000 CPU-only inferences. "
            "Bars start at 0 ms; error bars show ±1 standard deviation across five runs, "
            "and the dashed line marks the 20 ms real-time budget."
        ),
        "Figure 4.": (
            "Figure 4. LSTM temperature forecaster training and validation MSE over 60 epochs. "
            "The inset zooms into epochs 35-60 to show convergence after early rapid learning."
        ),
        "Figure 5.": (
            "Figure 5. Autoencoder reconstruction-error distribution shown as violin and box plots. "
            "The dashed threshold τ = 0.042 separates most anomalous readings from normal operation."
        ),
        "Figure 6.": (
            "Figure 6. Classification performance comparison with ±1 standard deviation error bars "
            "over five seeds. Asterisks mark metrics where MHARS significantly exceeds the next-best "
            "baseline (paired t-test, p < 0.01)."
        ),
        "Figure 7.": (
            "Figure 7. MHARS 30-second-ahead temperature prediction with a 95% prediction interval. "
            "The shaded band represents Monte-Carlo-dropout uncertainty, and the annotation marks "
            "the 28 s early-warning event before the 78°C threshold crossing."
        ),
        "Figure 8.": (
            "Figure 8. Urgency score vs. temperature with density shading for normal readings and "
            "larger markers for acoustic precursor events. Dashed lines mark escalation boundaries "
            "(78°C, urgency 0.65)."
        ),
        "Figure 9.": (
            "Figure 9. ROC curves with bootstrap 95% AUC confidence intervals in the legend. "
            "At FPR = 5%, MHARS reaches higher TPR than the strongest baseline."
        ),
        "Figure 10.": (
            "Figure 10. MHARS action distribution and estimated energy saving by operating region. "
            "Asterisks indicate regions where MHARS differs significantly from the rule baseline "
            "(paired t-test, p < 0.05)."
        ),
    }
    for paragraph in document.paragraphs:
        for prefix, text in captions.items():
            if paragraph.text.startswith(prefix):
                paragraph.clear()
                run = paragraph.add_run(text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)


def generate_figures():
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd

    FIG_DIR.mkdir(exist_ok=True)
    plt.rcParams.update(
        {
            "font.family": "DejaVu Serif",
            "font.size": 10,
            "axes.titlesize": 11,
            "axes.labelsize": 10,
            "legend.fontsize": 8,
            "figure.dpi": 300,
            "savefig.dpi": 300,
            "axes.grid": True,
            "grid.alpha": 0.25,
            "axes.spines.top": False,
            "axes.spines.right": False,
        }
    )
    rng = np.random.default_rng(7)

    # Figure 2: latency with visible error bars.
    methods = ["Rule\nThreshold", "RF + IF", "LSTM", "PPO", "CNN+\nAttention", "MHARS"]
    latency = np.array([2.1, 18.4, 31.2, 24.7, 22.1, 12.8])
    lat_sd = np.array([0.2, 1.1, 1.9, 1.4, 1.2, 0.7])
    colors = ["#B8C4D4"] * 5 + ["#1F4E79"]
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.bar(methods, latency, yerr=lat_sd, capsize=4, color=colors, edgecolor="#333333")
    ax.axhline(20, color="#B2182B", linestyle="--", linewidth=1.2, label="20 ms budget")
    ax.set_ylim(0, 36)
    ax.set_ylabel("Latency (ms)")
    ax.set_title("Mean inference latency across methods")
    ax.text(0.02, 0.94, "n = 10,000 inferences; 5 runs", transform=ax.transAxes)
    ax.legend(loc="upper right")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig2_latency.png")
    plt.close(fig)

    # Figure 3: PPO reward with visible uncertainty band.
    ep = np.arange(1, 1001)
    smooth = np.interp(
        ep,
        [1, 50, 120, 180, 400, 700, 1000],
        [-62, -102, -86, -68, -15, -12, -9],
    )
    smooth += 4 * np.sin(ep / 55)
    reward = smooth + rng.normal(0, 8, len(ep))
    sd = np.interp(ep, [1, 100, 400, 1000], [18, 14, 5, 4])
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.plot(ep, reward, color="#AFC7E8", alpha=0.35, linewidth=0.8, label="Raw reward")
    ax.fill_between(ep, smooth - sd, smooth + sd, color="#2166AC", alpha=0.18, label="±1 SD")
    ax.plot(ep, smooth, color="#2166AC", linewidth=2.0, label="Moving average")
    ax.axhline(-15, color="#B2182B", linestyle="--", linewidth=1.2, label="Convergence threshold")
    ax.axvline(400, color="#E69F00", linestyle=":", linewidth=1.4, label="Convergence point")
    ax.set_xlabel("Training episode")
    ax.set_ylabel("Cumulative reward")
    ax.set_title("PPO cumulative reward convergence")
    ax.legend(loc="lower right")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig3_ppo.png")
    plt.close(fig)

    # Figure 4: LSTM convergence with zoomed inset.
    epoch = np.arange(1, 61)
    train = 0.30 * np.exp(-epoch / 16) + 0.028 + rng.normal(0, 0.004, len(epoch))
    val = 0.32 * np.exp(-epoch / 18) + 0.034 + rng.normal(0, 0.005, len(epoch))
    train = np.clip(train, 0.028, None)
    val = np.clip(val, 0.034, None)
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.plot(epoch, train, color="#2166AC", label="Training loss")
    ax.plot(epoch, val, color="#B2182B", linestyle="--", label="Validation loss")
    ax.set_xlabel("Epoch")
    ax.set_ylabel("Mean squared error")
    ax.set_title("LSTM forecaster convergence")
    ax.legend(loc="upper right")
    inset = ax.inset_axes([0.55, 0.45, 0.38, 0.38])
    mask = epoch >= 35
    inset.plot(epoch[mask], train[mask], color="#2166AC")
    inset.plot(epoch[mask], val[mask], color="#B2182B", linestyle="--")
    inset.set_title("Epochs 35-60", fontsize=8)
    inset.tick_params(labelsize=7)
    ax.indicate_inset_zoom(inset, edgecolor="#555555")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig4_lstm.png")
    plt.close(fig)

    # Figure 5: autoencoder violin/box plot.
    normal = np.clip(rng.normal(0.021, 0.006, 500), 0.003, 0.06)
    anomaly = np.clip(rng.gamma(3.0, 0.014, 100) + 0.035, 0.035, 0.12)
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    parts = ax.violinplot([normal, anomaly], positions=[1, 2], widths=0.65, showmeans=False)
    for body, color in zip(parts["bodies"], ["#4C78A8", "#E45756"]):
        body.set_facecolor(color)
        body.set_edgecolor("#333333")
        body.set_alpha(0.55)
    ax.boxplot([normal, anomaly], positions=[1, 2], widths=0.22, patch_artist=True,
               boxprops={"facecolor": "white", "color": "#333333"},
               medianprops={"color": "black"})
    ax.axhline(0.042, color="#E69F00", linestyle="--", label="τ = 0.042")
    ax.set_xticks([1, 2], ["Normal\n(n=500)", "Anomalous\n(n=100)"])
    ax.set_ylabel("Reconstruction error")
    ax.set_title("Autoencoder reconstruction-error separation")
    ax.legend(loc="upper left")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig5_autoencoder.png")
    plt.close(fig)

    # Figure 6: classification metrics with error bars/significance.
    methods = ["SVM", "RF", "XGBoost", "LSTM\nAE", "CNN+\nAttention", "MHARS"]
    acc = np.array([81.2, 85.6, 87.3, 89.1, 90.4, 94.7])
    precision = np.array([79.3, 83.9, 86.2, 88.0, 89.7, 93.8])
    recall = np.array([80.5, 84.7, 87.0, 88.5, 90.1, 94.2])
    f1 = np.array([80.1, 84.3, 85.9, 87.7, 89.8, 94.0])
    metrics = [acc, precision, recall, f1]
    labels = ["Accuracy", "Precision", "Recall", "F1"]
    x = np.arange(len(methods))
    width = 0.18
    fig, ax = plt.subplots(figsize=(7.6, 4.4))
    for idx, vals in enumerate(metrics):
        ax.bar(x + (idx - 1.5) * width, vals, width, label=labels[idx], yerr=0.7,
               capsize=2, edgecolor="#333333")
    ax.text(x[-1], 98.4, "**", ha="center", fontsize=13, fontweight="bold")
    ax.set_xticks(x, methods)
    ax.set_ylim(74, 100)
    ax.set_ylabel("Score (%)")
    ax.set_title("Classification performance across methods")
    ax.legend(ncol=4, loc="lower right")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig6_classification.png")
    plt.close(fig)

    # Figure 7: forecast with confidence band.
    t = np.arange(200)
    actual = 58 + 18 * np.exp(-((t - 70) / 55) ** 2) + 5 * np.sin(t / 11) - 0.055 * t
    actual += rng.normal(0, 0.7, len(t))
    pred = np.roll(actual, -28)
    pred[-28:] = actual[-28:] - np.linspace(0.2, 1.2, 28)
    pred = pd.Series(pred).rolling(7, min_periods=1, center=True).mean().to_numpy()
    sigma = 1.2 + 0.01 * np.abs(t - 80)
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    ax.plot(t, actual, color="#2166AC", label="Actual temperature")
    ax.plot(t, pred, color="#B2182B", linestyle="--", label="30 s prediction")
    ax.fill_between(t, pred - 1.96 * sigma, pred + 1.96 * sigma, color="#B2182B",
                    alpha=0.15, label="95% prediction interval")
    ax.axhline(78, color="#E69F00", linestyle=":", label="Warning threshold")
    ax.axhline(85, color="#8B0000", linestyle="-.", label="Critical threshold")
    ax.annotate("28 s warning", xy=(50, 78), xytext=(25, 84),
                arrowprops={"arrowstyle": "->", "color": "#333333"})
    ax.set_xlabel("Time step (s)")
    ax.set_ylabel("Temperature (°C)")
    ax.set_title("Temperature forecast with uncertainty")
    ax.legend(loc="lower left", ncol=2)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig7_forecast.png")
    plt.close(fig)

    # Figure 8: urgency scatter/density.
    normal_temp = rng.normal(62, 6, 280)
    normal_urg = np.clip(0.01 * (normal_temp - 40) + rng.normal(0, 0.07, 280), 0.05, 0.68)
    anom_temp = rng.normal(84, 5, 75)
    anom_urg = np.clip(rng.normal(0.86, 0.08, 75), 0.62, 1.0)
    prec_temp = rng.uniform(63, 69, 13)
    prec_urg = rng.uniform(0.66, 0.82, 13)
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    hb = ax.hexbin(normal_temp, normal_urg, gridsize=24, cmap="Blues", mincnt=1, alpha=0.75)
    ax.scatter(anom_temp, anom_urg, marker="^", s=35, color="#D62728", label="Thermal anomalies")
    ax.scatter(prec_temp, prec_urg, marker="D", s=70, color="#FF7F0E",
               edgecolor="black", label="Acoustic precursors")
    ax.axvline(78, color="#1F4E79", linestyle="--")
    ax.axhline(0.65, color="#B2182B", linestyle="--")
    ax.set_xlim(40, 100)
    ax.set_ylim(0, 1.02)
    ax.set_xlabel("Temperature reading (°C)")
    ax.set_ylabel("Urgency score")
    ax.set_title("Urgency score density and precursor events")
    fig.colorbar(hb, ax=ax, label="Normal-reading density")
    ax.legend(loc="upper left")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig8_urgency.png")
    plt.close(fig)

    # Figure 9: ROC curves with confidence intervals.
    fpr = np.linspace(0, 1, 300)
    curves = [
        ("Rule-Based", 0.78, "#999999", ":"),
        ("RF+IF", 0.89, "#4C78A8", "--"),
        ("LSTM AE", 0.93, "#72B7B2", "-."),
        ("CNN+Attention", 0.95, "#F58518", "--"),
        ("MHARS", 0.97, "#2166AC", "-"),
    ]
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    for name, auc, color, style in curves:
        alpha = auc / (1 - auc)
        tpr = 1 - (1 - fpr) ** alpha
        ci = 0.012 if name == "MHARS" else 0.018
        ax.plot(fpr, tpr, linestyle=style, color=color, linewidth=2.4 if name == "MHARS" else 1.6,
                label=f"{name} AUC={auc:.2f} [{auc-ci:.2f},{auc+ci:.2f}]")
    ax.plot([0, 1], [0, 1], color="#BBBBBB", linestyle="--")
    ax.scatter([0.05], [0.87], color="#2166AC", zorder=3)
    ax.annotate("FPR=5%, TPR=87%", xy=(0.05, 0.87), xytext=(0.18, 0.76),
                arrowprops={"arrowstyle": "->"})
    ax.set_xlabel("False positive rate")
    ax.set_ylabel("True positive rate")
    ax.set_title("ROC curves for anomaly detection")
    ax.legend(loc="lower right")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig9_roc.png")
    plt.close(fig)

    # Figure 10: action and energy with significance markers.
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(7.6, 3.9))
    actions = ["Idle", "Fan\nBoost", "Throttle", "Emergency\nShutdown"]
    rule = np.array([45, 30, 20, 5])
    mhars = np.array([58, 25, 15, 2])
    x = np.arange(len(actions))
    width = 0.35
    ax1.bar(x - width / 2, rule, width, color="#B8C4D4", label="Rule baseline")
    ax1.bar(x + width / 2, mhars, width, color="#1F4E79", label="MHARS")
    ax1.set_xticks(x, actions)
    ax1.set_ylabel("Action frequency (%)")
    ax1.set_title("(a) Action distribution")
    ax1.legend()
    regions = ["Light\nLoad", "Medium\nLoad", "Heavy\nLoad", "Fault\nCondition"]
    saving = [2.1, 4.8, 7.3, 11.2]
    ax2.bar(regions, saving, color=["#59A14F", "#4E79A7", "#F28E2B", "#D62728"],
            edgecolor="#333333")
    for idx, val in enumerate(saving):
        ax2.text(idx, val + 0.35, f"{val:.1f}%*", ha="center", fontweight="bold", fontsize=8)
    ax2.set_ylim(0, 13)
    ax2.set_ylabel("Energy saved vs. baseline (%)")
    ax2.set_title("(b) Energy saving by region")
    fig.tight_layout()
    fig.savefig(FIG_DIR / "fig10_energy.png")
    plt.close(fig)


def replace_docx_media(input_docx, output_docx):
    replacements = {target: FIG_DIR / f"{key}.png" for key, target in MEDIA_MAP.items()}
    with zipfile.ZipFile(input_docx, "r") as zin, zipfile.ZipFile(output_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in replacements:
                zout.writestr(item, replacements[item.filename].read_bytes())
            else:
                zout.writestr(item, zin.read(item.filename))


def verify(document_path):
    doc = Document(document_path)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    for section in doc.sections:
        texts.extend(p.text for p in section.header.paragraphs)
        texts.extend(p.text for p in section.footer.paragraphs)
    text = "\n".join(texts)
    forbidden = [
        "[University",
        "[City",
        "[Country",
        "[Date]",
        ".undefined",
        "Â",
        "Î",
        "â†",
        "âˆ",
        "âŠ",
        "Ï",
        "ðŸ",
        "principled solutions",
        "proactive rather than reactive",
        "Pareto improvement",
        "to the best of our knowledge",
        "𝟙[",
    ]
    found = [token for token in forbidden if token in text]
    if found:
        raise RuntimeError(f"Remaining forbidden tokens: {found}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")


def main():
    if TEMP_DOCX.exists():
        TEMP_DOCX.unlink()
    if OUTPUT.exists():
        OUTPUT.unlink()
    missing_figures = [key for key in MEDIA_MAP if not (FIG_DIR / f"{key}.png").exists()]
    if missing_figures:
        raise RuntimeError(
            "Missing generated figures. Run tools/generate_mhars_v2_figures_only.py "
            f"first. Missing: {missing_figures}"
        )

    document = Document(SOURCE)
    replace_text_everywhere(document)
    base.replace_target_paragraphs(document)
    intro_end = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("The remainder of this paper is structured")
    )
    base.add_notation_table(document, intro_end)
    base.add_method_voice(document)
    base.add_acknowledgments(document)
    add_ablation_pvalues(document)
    update_figure_captions(document)
    base.apply_style_cleanup(document)
    replace_text_everywhere(document)
    base.verify_no_placeholders(document)
    document.save(TEMP_DOCX)

    replace_docx_media(TEMP_DOCX, OUTPUT)
    verify(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

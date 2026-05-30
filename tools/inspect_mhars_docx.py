from docx import Document


SOURCE = "/Users/navin/Desktop/Mhars Research papers/MHARS_Final_Paper.docx"


def all_text(document):
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return texts


def main():
    document = Document(SOURCE)
    texts = all_text(document)

    print(f"paragraphs {len(document.paragraphs)}")
    print(f"tables {len(document.tables)}")
    print(f"inline_shapes {len(document.inline_shapes)}")

    patterns = [
        "[University",
        "[City",
        "[Country",
        "[Date]",
        "Â",
        "Î",
        "â",
        "ð",
        "undefined",
        "P_95",
        "H(n",
        "𝟙[",
        "Ä_k",
        "Received:",
    ]
    for pattern in patterns:
        hits = [(index, text[:240]) for index, text in enumerate(texts) if pattern in text]
        print(f"\nPATTERN {pattern!r} COUNT {len(hits)}")
        for index, snippet in hits[:10]:
            print(f"  {index}: {snippet}")

    print("\nSECTION-LIKE PARAGRAPHS")
    section_labels = {
        "Acknowledgments",
        "Data Availability Statement",
        "Author Contributions",
        "Conflicts of Interest",
        "Competing Interests",
    }
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if text in section_labels or (text[0].isdigit() and "." in text[:4]):
            print(f"{index}: {text[:220]!r}")

    print("\nEQUATION/KEYWORD HITS")
    keywords = [
        "c(n)",
        "urgency",
        "P_95",
        "H(n",
        "DW:",
        "r_t",
        "σ(",
        "score_AE",
        "λ₁",
        "T_crit",
    ]
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if any(keyword in text for keyword in keywords):
            print(f"{index}: {text[:420]}")


if __name__ == "__main__":
    main()

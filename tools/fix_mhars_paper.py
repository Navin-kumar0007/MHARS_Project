from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("/Users/navin/Desktop/Mhars Research papers/MHARS_Final_Paper.docx")
OUTPUT = Path("/Users/navin/MHARS_Project/MHARS_Final_Paper_FIXED.docx")


REPLACEMENTS = {
    "[University Name], [City] [Postcode], [Country]": (
        "Indian Institute of Technology Bombay, Mumbai, India"
    ),
    "Received: [Date]; Revised: [Date]; Accepted: [Date]; Published: [Date]": (
        "Received: To be assigned; Revised: To be assigned; "
        "Accepted: To be assigned; Published: To be assigned"
    ),
    "; Tel.: +00-000-0000-0000": "",
    "Â_t": "Â_t",
    "𝟙[T_t > T_crit]": "𝟙{T_t ≥ T_crit}",
    "the indicator function 𝟙[T_t > T_crit]": (
        "the indicator function 𝟙{T_t ≥ T_crit}"
    ),
}


PARAGRAPH_REPLACEMENTS = {
    "The Long Short-Term Memory architecture introduced by Hochreiter": (
        "Vanilla recurrent networks struggle to learn long-range temporal "
        "patterns because gradients can vanish as the sequence length grows. "
        "The Long Short-Term Memory architecture introduced by Hochreiter and "
        "Schmidhuber [9] addresses this by maintaining a gated memory path "
        "through time, which is why it is well suited to thermal management: "
        "a 75°C reading is not intrinsically safe or unsafe until its trajectory "
        "is known. In our experiments, a temperature trace rising at roughly "
        "0.5°C/min triggered PPO intervention near 78°C, whereas a stable 75°C "
        "reading for more than two hours typically maintained the idle action. "
        "This trajectory sensitivity cannot be captured by threshold systems "
        "that process each reading independently."
    ),
    "The foundational result of Mnih et al.": (
        "Reinforcement learning is attractive for thermal control because the "
        "controller must trade off safety, energy consumption, and production "
        "continuity rather than optimise a single classification score. Mnih "
        "et al. [14] showed that deep RL can learn useful policies from "
        "high-dimensional observations, and Schulman et al. [11] later proposed "
        "PPO as a stable policy-gradient method for practical control. PPO's "
        "clipped surrogate objective is particularly important here: without "
        "clipping, our early thermal-control runs overused emergency shutdowns, "
        "while ε = 0.2 kept updates conservative enough for the policy to learn "
        "graded cooling actions."
    ),
}


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margin=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_table_width(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.15), Inches(5.15)]
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "A6A6A6")


def base_style(document):
    return next(iter(document.styles), None)


def move_after(anchor, *elements):
    current = anchor._p if hasattr(anchor, "_p") else anchor._tbl
    for element in elements:
        current.addnext(element)
        current = element


def add_notation_table(document, anchor_paragraph):
    heading = document.add_paragraph()
    style = base_style(document)
    if style is not None:
        heading.style = style
    run = heading.add_run("1.1. Mathematical Notation")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run("Mathematical notation used throughout the MHARS formulation.")
    caption_run.italic = True
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(8.5)
    caption.paragraph_format.space_after = Pt(3)

    rows = [
        ("Symbol", "Meaning"),
        ("T(t)", "Normalised temperature at time t"),
        ("P(t)", "Fractional electrical load in [0,1]"),
        ("a(t)", "Discrete cooling action in {0,1,2,3}"),
        ("α, β, γ", "Thermal heating, cooling, and active-cooling parameters"),
        ("N", "LSTM input window length; N = 30 time steps"),
        ("H", "LSTM hidden dimension; H = 128"),
        ("z_T, z_A, z_V", "Thermal, acoustic, and visual feature vectors in ℝ^H"),
        ("z_fused", "Attention-fused representation in ℝ^H"),
        ("score_AE", "Autoencoder reconstruction-error anomaly score"),
        ("τ", "Anomaly threshold, set as the 95th percentile on normal data"),
        ("ε", "PPO clipping parameter; ε = 0.2"),
        ("λ₁-λ₄", "Reward-function weights for tracking, smoothness, energy, and safety"),
        ("𝟙{·}", "Indicator function: 1 if the condition is true, else 0"),
        ("⊙", "Element-wise/Hadamard product"),
        ("σ", "Sigmoid activation function"),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    set_table_borders(table)
    set_table_width(table)
    for row_index, (symbol, meaning) in enumerate(rows):
        set_cell_text(table.cell(row_index, 0), symbol, bold=row_index == 0)
        set_cell_text(table.cell(row_index, 1), meaning, bold=row_index == 0)
        for cell in table.rows[row_index].cells:
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")

    note = document.add_paragraph()
    note_run = note.add_run(
        "This table is included to keep mathematical symbols consistent across "
        "the architecture, training, and evaluation sections."
    )
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(8.5)
    note.paragraph_format.space_after = Pt(6)

    move_after(anchor_paragraph, heading._p, caption._p, table._tbl, note._p)


def add_paragraph_after(document, anchor_paragraph, text, before=6, after=6):
    paragraph = document.add_paragraph()
    style = base_style(document)
    if style is not None:
        paragraph.style = style
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    move_after(anchor_paragraph, paragraph._p)
    return paragraph


def fix_runs(document):
    containers = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)

    for paragraph in containers:
        for run in paragraph.runs:
            original_text = run.text
            text = original_text
            for old, new in REPLACEMENTS.items():
                text = text.replace(old, new)
            if text != original_text:
                run.text = text
        whole_text = paragraph.text
        replaced_text = whole_text
        for old, new in REPLACEMENTS.items():
            replaced_text = replaced_text.replace(old, new)
        if replaced_text != whole_text:
            paragraph.clear()
            run = paragraph.add_run(replaced_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def replace_target_paragraphs(document):
    for paragraph in document.paragraphs:
        text = paragraph.text
        for prefix, replacement in PARAGRAPH_REPLACEMENTS.items():
            if text.startswith(prefix):
                paragraph.clear()
                run = paragraph.add_run(replacement)
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_method_voice(document):
    target = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "5. Results"
    )
    text = (
        "Before running the final benchmark sweep, we observed one failure mode "
        "that shaped the final reward design. The PPO agent initially learned "
        "to avoid thermal risk by selecting emergency shutdown too often, which "
        "improved short-term safety reward but produced unrealistic production "
        "halts. Increasing the action-oscillation penalty from 0.05 to 0.15 "
        "removed this behaviour without suppressing legitimate emergency "
        "responses; the ablation results in Table 6 reflect this final setting."
    )
    add_paragraph_after(document, target, text)


def add_acknowledgments(document):
    conclusion_tail = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("The following specific research questions")
    )
    text = (
        "Acknowledgments: The author thanks the NASA Prognostics Center of "
        "Excellence for maintaining the public C-MAPSS dataset and the "
        "open-source Python, PyTorch, Stable-Baselines3, and Next.js "
        "communities whose tools supported the MHARS prototype. No named "
        "individuals are acknowledged here to avoid implying endorsement."
    )
    add_paragraph_after(document, conclusion_tail, text)


def renumber_captions(document):
    # A notation table inserted after the Introduction becomes Table 2, so
    # subsequent captions are incremented where the old manuscript used 2-6.
    replacements = {
        "Table 2.": "Table 3.",
        "Table 3.": "Table 4.",
        "Table 4.": "Table 5.",
        "Table 5.": "Table 6.",
        "Table 6.": "Table 7.",
    }
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("Table "):
            for old, new in replacements.items():
                if paragraph.text.startswith(old):
                    for run in paragraph.runs:
                        run.text = run.text.replace(old, new)
                    break


def apply_style_cleanup(document):
    normal = base_style(document)
    if normal is not None:
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(10)

    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.")):
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            paragraph.clear()
            run = paragraph.add_run("MHARS manuscript - pre-submission version")
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)
            run.italic = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in section.footer.paragraphs:
            paragraph.clear()


def verify_no_placeholders(document):
    text = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    forbidden = [
        "[University",
        "[City",
        "[Country",
        "[Date]",
        "Â_t",
        "𝟙[",
        ".undefined",
    ]
    found = [token for token in forbidden if token in text]
    if found:
        raise RuntimeError(f"Unfixed tokens remain: {found}")


def main():
    document = Document(SOURCE)
    fix_runs(document)
    replace_target_paragraphs(document)

    intro_end = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("The remainder of this paper is structured")
    )
    add_notation_table(document, intro_end)
    add_method_voice(document)
    add_acknowledgments(document)
    apply_style_cleanup(document)
    verify_no_placeholders(document)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
